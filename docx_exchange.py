from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt



doc = Document(docx="сведения 2 таблицы.docx")
doc_2 = Document(docx="Приказ о создании комиссии.docx")

def change_table_text(dict):
    table = doc.tables[0]

    table.rows[2].cells[2].text = f"{dict.get('Наименование объекта')}"
    table.rows[3].cells[2].text = f"{dict.get('Адрес размещения')}"
    table.rows[4].cells[2].text = f"{dict.get('Сфера деятельности')}"
    table.rows[5].cells[2].text = f"{dict.get('Назначение объекта')}"
    table.rows[6].cells[2].text = f"{dict.get('Тип объекта')}"
    table.rows[7].cells[2].text = f"{dict.get('Архитектура объекта')}"
    table.rows[8].cells[2].text = f"{dict.get('Наименование ЭВМ')}"
    table.rows[9].cells[2].text = f"{dict.get('Наименование ПО')}"
    table.rows[10].cells[2].text = f"{dict.get('Наименование прикладных ПО')}"
    table.rows[11].cells[2].text = f"{dict.get('Категория сети электросвязи')}"
    table.rows[13].cells[2].text = f"{dict.get('Наименование оператора связи')}"
    table.rows[14].cells[2].text = f"{dict.get('Основные угрозы безопасности')}"
    table.rows[15].cells[2].text = f"{dict.get('Типы компьютерных инцидентов')}"
    table.rows[17].cells[2].text = f"{dict.get('Реализованные меры защиты')}"
    table.rows[18].cells[2].text = f"{dict.get('Применяемые средства защиты')}"
    table.rows[20].cells[2].text = f"{dict.get('Категория нарушителя')}"
    table.rows[22].cells[2].text = f"{dict.get('Организационные меры')}"

    doc.save("Таблица категорирования.docx")


def create_order(ex_word, new_word, par):
    docx = doc_2
    ex_word = ex_word
    par = int(par)

    for i in docx.paragraphs[par].runs:
        if i.text == ex_word:
            i.text = i.text.replace(i.text, new_word)

    docx.save("New_order.docx")


def order_change():
    create_order(ex_word="Москва", new_word="Казань", par=1)
    create_order(ex_word="num", new_word="555", par=0)
    create_order(ex_word="num", new_word="19", par=2)
    create_order(ex_word="february", new_word="января", par=2)
    create_order(ex_word="nm", new_word="23", par=2)
    create_order(ex_word="vimpel", new_word="ОАО Согаз", par=12)
    create_order(ex_word="vimpel", new_word="ОАО Согаз", par=13)
    create_order(ex_word="директор", new_word="Генеральный директор", par=30)
    create_order(ex_word="vimpel", new_word="ОАО Согаз", par=30)
    create_order(ex_word="Петров", new_word="С.С.Петров", par=31)

def order_table():
    table = doc_2.tables[0]
    table_element = table._tbl
    header = "Хатыпов Р.Р."
    comission = ["Черных В.А.", "Заикин Р.В.", "Белкин А.С."]
    if header != "":
        table.rows[0].cells[1].paragraphs[0].add_run(header)
        table.rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(len(comission)):
            table.rows[i+1].cells[1].paragraphs[0].add_run(comission[i])
            table.rows[i+1].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i in range(10, len(comission), -1):
            RowA = table.rows[i]
            row_element = RowA._tr
            table_element.remove(row_element)


    doc_2.save("New_order.docx")


if __name__ == "__main__":
    order_change()
    order_table()